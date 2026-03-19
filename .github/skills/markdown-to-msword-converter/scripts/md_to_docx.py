"""Markdown -> Word (.docx) converter.

Converts a single Markdown file into a Word document using `python-docx`.
Includes best-effort support for headings, lists, tables, quotes, and hyperlinks.
"""

import argparse
import os
import re
from pathlib import Path
from urllib.parse import quote, unquote


LINK_PATTERN = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def strip_inline_markdown(text: str) -> str:
    """Remove common inline markdown formatting, including links."""
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    return text


def strip_text_markdown_only(text: str) -> str:
    """Remove inline markdown formatting but preserve link structure elsewhere."""
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text


def resolve_word_link_target(source_md: Path, target: str) -> str:
    """Resolve a Markdown link target to a Word-friendly external hyperlink target.

    - Absolute URLs are returned unchanged.
    - Fragment-only links are returned unchanged.
    - Relative paths are resolved against `source_md` parent.
    - Internal file links are converted to relative targets.
    - `.md` targets are rewritten to `.docx`.
    """
    cleaned = target.strip().strip("<>")
    if not cleaned:
        return target

    if re.match(r"^[a-zA-Z][a-zA-Z0-9+.-]*:", cleaned):
        return cleaned

    if cleaned.startswith("#"):
        return cleaned

    path_part, frag = (cleaned.split("#", 1) + [""])[:2]
    path_part = path_part.split("?", 1)[0]
    if not path_part:
        return cleaned

    resolved = (source_md.parent / unquote(path_part)).resolve()

    if resolved.exists() and resolved.is_dir():
        directory_candidates = [
            resolved / "index.docx",
            resolved / "README.docx",
            (resolved / "index.md").with_suffix(".docx"),
            (resolved / "README.md").with_suffix(".docx"),
        ]
        for candidate in directory_candidates:
            if candidate.exists() and candidate.is_file():
                resolved = candidate
                break

    if resolved.exists() and resolved.is_dir():
        markdown_candidates = [
            resolved / "Overview.md",
            resolved / "README.md",
            resolved / "index.md",
        ]

        top_level_markdowns = sorted(resolved.glob("*.md"))
        recursive_markdowns = sorted(resolved.rglob("*.md"))

        for candidate in markdown_candidates + top_level_markdowns + recursive_markdowns:
            if candidate.exists() and candidate.is_file():
                resolved = candidate.with_suffix(".docx")
                break

    if resolved.exists() and resolved.is_dir():
        preferred_exts = [".docx", ".xlsx", ".xls", ".pdf", ".pptx", ".ppt", ".csv", ".txt"]
        top_level_files = sorted(p for p in resolved.iterdir() if p.is_file())
        recursive_files = sorted(p for p in resolved.rglob("*") if p.is_file())

        for ext in preferred_exts:
            for candidate in top_level_files + recursive_files:
                if candidate.suffix.lower() == ext:
                    resolved = candidate
                    break
            if not resolved.is_dir():
                break

    if resolved.suffix.lower() == ".md":
        resolved = resolved.with_suffix(".docx")

    base_dir = source_md.parent.resolve()
    relative_target = resolved.relative_to(base_dir) if resolved.is_relative_to(base_dir) else Path(os.path.relpath(str(resolved), str(base_dir)))

    # Word handles external hyperlink targets best as URI-like forward-slash paths.
    target_path = relative_target.as_posix()
    # Avoid Word coercing bare filenames into https:// links by making local
    # relative file targets explicit.
    if not target_path.startswith("../") and not target_path.startswith("./"):
        target_path = f"./{target_path}"

    target_path = quote(target_path, safe="/-._~")
    return f"{target_path}#{frag}" if frag else target_path


def normalize_link_label(label: str) -> str:
    """Normalize visible link text for Word output."""
    stripped = label.strip()
    if stripped.lower().endswith(".md"):
        return stripped[:-3]
    return stripped


def add_hyperlink(paragraph, text: str, url: str) -> None:
    """Add an external hyperlink run to a paragraph.

    Uses python-docx's underlying OOXML APIs to create a relationship and
    append a `w:hyperlink` element.
    """
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    # Be explicit about the relationship namespace to avoid prefix/namespace
    # edge cases across python-docx versions.
    hyperlink.set("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id", r_id)
    hyperlink.set(qn("w:history"), "1")

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    new_run.append(r_pr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_markdown_runs(paragraph, text: str, source_md: Path) -> None:
    last = 0
    for match in LINK_PATTERN.finditer(text):
        before = text[last:match.start()]
        if before:
            paragraph.add_run(strip_text_markdown_only(before))

        label = normalize_link_label(strip_text_markdown_only(match.group(1).strip()))
        target = match.group(2).strip()
        if target:
            resolved_target = resolve_word_link_target(source_md, target)
            add_hyperlink(paragraph, label or target, resolved_target)
        else:
            paragraph.add_run(label)

        last = match.end()

    trailing = text[last:]
    if trailing:
        paragraph.add_run(strip_text_markdown_only(trailing))


def is_table_line(line: str) -> bool:
    stripped = line.strip()
    return "|" in stripped and stripped.startswith("|") and stripped.endswith("|")


def parse_table(lines: list[str], start_idx: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start_idx

    while i < len(lines) and is_table_line(lines[i]):
        row = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
        rows.append(row)
        i += 1

    if len(rows) >= 2:
        sep = rows[1]
        if all(re.match(r"^:?-{3,}:?$", cell.replace(" ", "")) for cell in sep):
            rows.pop(1)

    return rows, i


def convert_markdown_to_docx(input_path: Path, output_path: Path) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Missing dependency 'python-docx'. Install with: python -m pip install python-docx") from exc

    lines = input_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
            i += 1
            continue

        if is_table_line(line):
            rows, next_i = parse_table(lines, i)
            if rows:
                col_count = max(len(r) for r in rows)
                table = doc.add_table(rows=1, cols=col_count)
                table.style = "Table Grid"

                for c in range(col_count):
                    table.rows[0].cells[c].text = strip_inline_markdown(rows[0][c] if c < len(rows[0]) else "")

                for row in rows[1:]:
                    row_cells = table.add_row().cells
                    for c in range(col_count):
                        row_cells[c].text = strip_inline_markdown(row[c] if c < len(row) else "")

                i = next_i
                continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            level = min(len(heading.group(1)), 4)
            doc.add_heading(strip_inline_markdown(heading.group(2).strip()), level=level)
            i += 1
            continue

        bullet = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet:
            paragraph = doc.add_paragraph("")
            add_markdown_runs(paragraph, bullet.group(1).strip(), input_path)
            paragraph.style = "List Bullet"
            i += 1
            continue

        numbered = re.match(r"^\d+[.)]\s+(.*)$", stripped)
        if numbered:
            paragraph = doc.add_paragraph("")
            add_markdown_runs(paragraph, numbered.group(1).strip(), input_path)
            paragraph.style = "List Number"
            i += 1
            continue

        quote = re.match(r"^>\s?(.*)$", stripped)
        if quote:
            paragraph = doc.add_paragraph("")
            add_markdown_runs(paragraph, quote.group(1).strip(), input_path)
            paragraph.style = "Intense Quote"
            i += 1
            continue

        paragraph = doc.add_paragraph("")
        add_markdown_runs(paragraph, stripped, input_path)
        i += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert a single Markdown (.md) file to Word (.docx)")
    parser.add_argument("input", help="Path to source Markdown file")
    parser.add_argument("--output", "-o", help="Path to output DOCX file")
    args = parser.parse_args()

    input_path = Path(args.input).expanduser().resolve()
    if not input_path.exists() or not input_path.is_file():
        raise FileNotFoundError(f"Input file not found: {input_path}")
    if input_path.suffix.lower() != ".md":
        raise ValueError("Input file must be a .md file")

    output_path = Path(args.output).expanduser().resolve() if args.output else input_path.with_suffix(".docx")

    convert_markdown_to_docx(input_path, output_path)
    print(f"Converted: {input_path}")
    print(f"Output:    {output_path}")


if __name__ == "__main__":
    main()
